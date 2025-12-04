"""
Document Loader

Multi-format document parsing and metadata extraction for knowledge base construction.

Supported Formats:
- Markdown (.md, .markdown)
- Plain Text (.txt)
- PDF (.pdf)
- Code (.py, .js, .java, .cpp, .c, .h, .rs, .go, etc.)
- JSON (.json)
- HTML (.html, .htm)
- DOCX (.docx) - optional

Author: Animation AI Studio
Date: 2025-12-03
"""

import logging
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
import mimetypes

from ..common import (
    Document,
    DocumentType,
    ProcessingStatus,
    generate_document_id
)

logger = logging.getLogger(__name__)


class DocumentLoader:
    """
    Multi-format document loader with metadata extraction

    Features:
    - Magic byte and extension-based type detection
    - Markdown parsing with header extraction
    - PDF text extraction (pdfplumber/PyPDF2)
    - Code file loading with syntax preservation
    - JSON structured data loading
    - HTML text extraction (BeautifulSoup)
    - Metadata extraction (title, author, timestamps)
    - Batch loading with progress tracking
    - Error handling and recovery
    """

    # Supported code file extensions
    CODE_EXTENSIONS = {
        '.py', '.pyw',  # Python
        '.js', '.jsx', '.ts', '.tsx',  # JavaScript/TypeScript
        '.java', '.kt',  # Java/Kotlin
        '.cpp', '.cc', '.cxx', '.hpp', '.h',  # C++
        '.c', '.h',  # C
        '.rs',  # Rust
        '.go',  # Go
        '.rb',  # Ruby
        '.php',  # PHP
        '.swift',  # Swift
        '.sh', '.bash',  # Shell
        '.sql',  # SQL
        '.r', '.R',  # R
        '.scala',  # Scala
        '.m', '.mm',  # Objective-C
        '.cs',  # C#
        '.lua',  # Lua
        '.pl', '.pm',  # Perl
    }

    def __init__(self, encoding: str = 'utf-8', errors: str = 'replace'):
        """
        Initialize document loader

        Args:
            encoding: Text encoding (default: utf-8)
            errors: Error handling for decoding (default: replace)
        """
        self.encoding = encoding
        self.errors = errors

        # Initialize mimetypes
        mimetypes.init()

    def load_document(self, file_path: Path) -> Document:
        """
        Load single document with automatic format detection

        Args:
            file_path: Path to document file

        Returns:
            Document object with content and metadata

        Raises:
            FileNotFoundError: If file doesn't exist
            RuntimeError: If loading fails
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if not file_path.is_file():
            raise RuntimeError(f"Path is not a file: {file_path}")

        try:
            # Detect document type
            doc_type = self._detect_type(file_path)

            # Load content based on type
            if doc_type == DocumentType.MARKDOWN:
                content = self._load_markdown(file_path)
            elif doc_type == DocumentType.PDF:
                content = self._load_pdf(file_path)
            elif doc_type == DocumentType.CODE:
                content = self._load_code(file_path)
            elif doc_type == DocumentType.JSON:
                content = self._load_json(file_path)
            elif doc_type == DocumentType.HTML:
                content = self._load_html(file_path)
            elif doc_type == DocumentType.DOCX:
                content = self._load_docx(file_path)
            elif doc_type == DocumentType.TEXT:
                content = self._load_text(file_path)
            else:
                # Fallback to text
                content = self._load_text(file_path)

            # Extract metadata
            metadata = self._extract_metadata(file_path, content, doc_type)

            # Get file stats
            stats = file_path.stat()

            # Create document
            document = Document(
                id=generate_document_id(file_path),
                path=file_path,
                doc_type=doc_type,
                content=content,
                metadata=metadata,
                created_at=stats.st_ctime,
                modified_at=stats.st_mtime,
                file_size_bytes=stats.st_size,
                processing_status=ProcessingStatus.COMPLETED
            )

            logger.info(f"Loaded document: {file_path.name} ({doc_type.value}, {len(content)} chars)")

            return document

        except Exception as e:
            logger.error(f"Failed to load document {file_path}: {e}")
            raise RuntimeError(f"Document loading failed: {e}") from e

    def batch_load(
        self,
        directory: Path,
        pattern: str = "**/*",
        recursive: bool = True,
        skip_errors: bool = True
    ) -> List[Document]:
        """
        Batch load documents from directory

        Args:
            directory: Directory to scan
            pattern: Glob pattern for files (default: **/*)
            recursive: Recursively scan subdirectories
            skip_errors: Skip files that fail to load

        Returns:
            List of loaded documents
        """
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")

        if not directory.is_dir():
            raise RuntimeError(f"Path is not a directory: {directory}")

        documents = []
        file_paths = list(directory.glob(pattern))

        logger.info(f"Found {len(file_paths)} files in {directory}")

        for file_path in file_paths:
            if not file_path.is_file():
                continue

            try:
                document = self.load_document(file_path)
                documents.append(document)
            except Exception as e:
                if skip_errors:
                    logger.warning(f"Skipping {file_path.name}: {e}")
                else:
                    raise

        logger.info(f"Loaded {len(documents)} documents from {directory}")

        return documents

    def _detect_type(self, file_path: Path) -> DocumentType:
        """
        Detect document type from extension and content

        Args:
            file_path: Path to file

        Returns:
            DocumentType enum value
        """
        extension = file_path.suffix.lower()

        # Extension-based detection
        if extension in {'.md', '.markdown'}:
            return DocumentType.MARKDOWN
        elif extension == '.pdf':
            return DocumentType.PDF
        elif extension in {'.json', '.jsonl'}:
            return DocumentType.JSON
        elif extension in {'.html', '.htm'}:
            return DocumentType.HTML
        elif extension == '.docx':
            return DocumentType.DOCX
        elif extension in self.CODE_EXTENSIONS:
            return DocumentType.CODE
        elif extension in {'.txt', '.text'}:
            return DocumentType.TEXT

        # Mimetype detection
        mime_type, _ = mimetypes.guess_type(str(file_path))
        if mime_type:
            if 'text' in mime_type:
                return DocumentType.TEXT
            elif 'json' in mime_type:
                return DocumentType.JSON
            elif 'html' in mime_type:
                return DocumentType.HTML

        # Default to text
        return DocumentType.TEXT

    def _load_text(self, file_path: Path) -> str:
        """Load plain text file"""
        return file_path.read_text(encoding=self.encoding, errors=self.errors)

    def _load_markdown(self, file_path: Path) -> str:
        """
        Load Markdown file

        Preserves formatting and structure
        """
        return file_path.read_text(encoding=self.encoding, errors=self.errors)

    def _load_pdf(self, file_path: Path) -> str:
        """
        Load PDF file using pdfplumber (preferred) or PyPDF2 (fallback)

        Returns:
            Extracted text content
        """
        try:
            # Try pdfplumber first (better quality)
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return "\n\n".join(text_parts)

        except ImportError:
            # Fallback to PyPDF2
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                text_parts = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return "\n\n".join(text_parts)

            except ImportError:
                raise RuntimeError("PDF support requires pdfplumber or PyPDF2: pip install pdfplumber")

    def _load_code(self, file_path: Path) -> str:
        """
        Load code file

        Preserves syntax and formatting
        """
        return file_path.read_text(encoding=self.encoding, errors=self.errors)

    def _load_json(self, file_path: Path) -> str:
        """
        Load JSON file

        Returns formatted JSON as string
        """
        import json
        with open(file_path, 'r', encoding=self.encoding) as f:
            data = json.load(f)
        return json.dumps(data, indent=2, ensure_ascii=False)

    def _load_html(self, file_path: Path) -> str:
        """
        Load HTML file and extract text

        Uses BeautifulSoup if available, otherwise returns raw HTML
        """
        html_content = file_path.read_text(encoding=self.encoding, errors=self.errors)

        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')

            # Remove script and style elements
            for script in soup(['script', 'style']):
                script.decompose()

            # Get text
            text = soup.get_text(separator='\n', strip=True)
            return text

        except ImportError:
            logger.warning("BeautifulSoup not available, returning raw HTML")
            # Basic HTML tag removal
            text = re.sub(r'<[^>]+>', '', html_content)
            return text

    def _load_docx(self, file_path: Path) -> str:
        """
        Load DOCX file

        Requires python-docx
        """
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs]
            return "\n\n".join(paragraphs)

        except ImportError:
            raise RuntimeError("DOCX support requires python-docx: pip install python-docx")

    def _extract_metadata(
        self,
        file_path: Path,
        content: str,
        doc_type: DocumentType
    ) -> Dict[str, Any]:
        """
        Extract metadata from document

        Args:
            file_path: Path to file
            content: Document content
            doc_type: Document type

        Returns:
            Metadata dictionary
        """
        metadata = {
            "filename": file_path.name,
            "extension": file_path.suffix,
            "doc_type": doc_type.value,
        }

        # Extract title
        title = self._extract_title(content, doc_type)
        if title:
            metadata["title"] = title

        # Markdown-specific metadata
        if doc_type == DocumentType.MARKDOWN:
            headers = self._extract_markdown_headers(content)
            if headers:
                metadata["headers"] = headers

        # Code-specific metadata
        if doc_type == DocumentType.CODE:
            metadata["language"] = file_path.suffix[1:]  # Remove leading dot

        return metadata

    def _extract_title(self, content: str, doc_type: DocumentType) -> Optional[str]:
        """
        Extract document title

        Args:
            content: Document content
            doc_type: Document type

        Returns:
            Title string or None
        """
        if doc_type == DocumentType.MARKDOWN:
            # Look for first # header
            match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
            if match:
                return match.group(1).strip()

        # Look for title in first few lines
        lines = content.split('\n')[:5]
        for line in lines:
            line = line.strip()
            if line and len(line) < 100:
                return line

        return None

    def _extract_markdown_headers(self, content: str) -> List[str]:
        """
        Extract all headers from Markdown

        Args:
            content: Markdown content

        Returns:
            List of header texts
        """
        headers = re.findall(r'^#+\s+(.+)$', content, re.MULTILINE)
        return headers
